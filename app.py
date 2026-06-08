from flask import Flask, request, jsonify
from docx import Document
import base64
import io
import pdfplumber
import re
import os

app = Flask(__name__)

# ------------------------
# PII MASKING FUNCTION
# ------------------------
def mask_pii(text: str) -> str:
    if not text:
        return text

    # 1️⃣ Emails
    text = re.sub(
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        '[EMAIL_REMOVED]',
        text
    )

    # 2️⃣ Phone numbers (international + local)
    text = re.sub(
        r'(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}',
        '[PHONE_REMOVED]',
        text
    )

    # 3️⃣ LinkedIn URLs
    text = re.sub(
        r'linkedin\.com\/[^\s]+',
        '[LINKEDIN_REMOVED]',
        text,
        flags=re.IGNORECASE
    )

    # 4️⃣ Other URLs
    text = re.sub(
        r'(http|https):\/\/[^\s]+',
        '[URL_REMOVED]',
        text
    )

    # 5️⃣ Remove name line (heuristic)
    lines = text.splitlines()
    cleaned_lines = []

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Heuristic: first non-empty line with 2–4 capitalized words
        if i < 3 and re.match(r'^([A-Z][a-z]+ ){1,3}[A-Z][a-z]+$', stripped):
            continue  # likely candidate name

        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)

    return text.strip()
    

def extract_docx_text(doc_bytes):

    doc_stream = io.BytesIO(doc_bytes)

    document = Document(doc_stream)

    text = []

    # paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)
    

@app.route("/extract", methods=["POST"])
def extract_text():

    data = request.get_json()
    if not data or "fileBase64" not in data:
        return jsonify({"error": "fileBase64 is required"}), 400

    try:
        # Decode Base64 → PDF bytes
        try:
            file_bytes = base64.b64decode(data["fileBase64"])
        except Exception:
            return jsonify({
                "success": False,
                "error": "Invalid Base64 data"
            }), 400

        file_name = data.get("fileName")

        if not file_name:
            return jsonify({
                "success": False,
                "error": "fileName is required"
            }), 400
        
        file_name = file_name.lower()
        if file_name.endswith(".pdf"):
            pdf_stream = io.BytesIO(file_bytes)
            extracted_text = ""
            with pdfplumber.open(pdf_stream) as pdf:
                for page in pdf.pages:
                    extracted_text += (
                        page.extract_text() or ""
                    ) + "\n"
                    
        elif file_name.endswith(".docx"):
            extracted_text = extract_docx_text(file_bytes)
        
        else:
            return jsonify({
                "success": False,
                "error": "Only PDF and DOCX files are supported"
            }), 400

        # 🔐 MASK PII HERE
        masked_text = mask_pii(extracted_text)

        return jsonify({
            "success": True,
            "fileType": file_name.split('.')[-1],
            "text": masked_text.strip()
        })

    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
